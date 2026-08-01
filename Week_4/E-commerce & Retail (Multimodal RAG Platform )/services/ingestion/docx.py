from docx import Document

def extract_docx(file_path: str) -> list[str]:
    """
    Extract text from DOCX paragraphs and tables.
    Returns a list of text strings.
    """

    document = Document(file_path)
    extracted_text = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            extracted_text.append(text)
    # Extract tables
    for table in document.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                extracted_text.append(" | ".join(row_text))

    return extracted_text